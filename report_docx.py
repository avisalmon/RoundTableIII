"""Phase III report: markdown source to an editable Word document.

The steering committee edits this file in Word with track changes, so the output is
built for editing rather than for looks. Real heading styles so the navigation pane
works, a TOC field Word rebuilds itself after the text moves, and tables that are
still tables. No images, no text boxes, no drawing canvases: everything here
survives a track-changes pass.

The website and the book take a different route entirely, through HTML and headless
Chrome. Only the small markdown helpers are shared with `build_book.py`.

    python report_docx.py
"""

from __future__ import annotations

import re
import sys
from dataclasses import dataclass, field
from pathlib import Path

import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from build_book import (
    column_alignments,
    parse_front_matter,
    smart_typography,
    split_table_row,
)

ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "report" / "phase3_report.md"
# Named for the web, because the site publishes this file and `build_book.py` links it.
OUTPUT = ROOT / "phase3-report.docx"

# A4 at 11pt over 2.5cm margins runs a little under 600 words to the page once
# headings and paragraph spacing are counted. Used only for the size report.
WORDS_PER_PAGE = 560
ROWS_PER_PAGE = 30

ALIGNMENTS = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
}


# --------------------------------------------------------------------------- parsing


@dataclass
class Block:
    kind: str
    lines: list[str] = field(default_factory=list)
    level: int = 0
    ordered: bool = False


def parse_blocks(body: str) -> list[Block]:
    blocks: list[Block] = []
    lines = body.splitlines()
    index = 0

    def flush(current: Block | None) -> None:
        if current and current.lines:
            blocks.append(current)

    paragraph: Block | None = None
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if not stripped:
            flush(paragraph)
            paragraph = None
            index += 1
            continue

        marker = re.match(r"^<!--\s*(pagebreak|toc)\s*-->$", stripped)
        if marker:
            flush(paragraph)
            paragraph = None
            blocks.append(Block(marker.group(1)))
            index += 1
            continue

        if stripped.startswith("<!--"):  # a note to the editor, never rendered
            index += 1
            continue

        heading = re.match(r"^(#{1,4})\s+(.+?)\s*$", stripped)
        if heading:
            flush(paragraph)
            paragraph = None
            blocks.append(Block("heading", [heading.group(2)], level=len(heading.group(1))))
            index += 1
            continue

        # A pipe table needs its divider row on the next line to count as one.
        if stripped.startswith("|") and index + 1 < len(lines) and re.match(r"^\s*\|[\s:|-]+\|\s*$", lines[index + 1]):
            flush(paragraph)
            paragraph = None
            table = Block("table")
            while index < len(lines) and lines[index].strip().startswith("|"):
                table.lines.append(lines[index])
                index += 1
            blocks.append(table)
            continue

        bullet = re.match(r"^[-*]\s+(.+)$", stripped)
        number = re.match(r"^\d+[.)]\s+(.+)$", stripped)
        if bullet or number:
            flush(paragraph)
            paragraph = None
            item = Block("list", ordered=bool(number))
            while index < len(lines):
                entry = lines[index].strip()
                nxt = re.match(r"^[-*]\s+(.+)$", entry) if item.ordered is False else re.match(r"^\d+[.)]\s+(.+)$", entry)
                if not nxt:
                    break
                item.lines.append(nxt.group(1))
                index += 1
            blocks.append(item)
            continue

        if stripped.startswith(">"):
            flush(paragraph)
            paragraph = None
            quote = Block("quote")
            while index < len(lines) and lines[index].strip().startswith(">"):
                quote.lines.append(lines[index].strip().lstrip(">").strip())
                index += 1
            blocks.append(quote)
            continue

        if paragraph is None:
            paragraph = Block("paragraph")
        paragraph.lines.append(stripped)
        index += 1

    flush(paragraph)
    return blocks


# --------------------------------------------------------------------------- runs

INLINE = re.compile(r"(\*\*[^*]+\*\*|(?<!\*)\*[^*\n]+\*(?!\*)|\[[^\]]+\]\([^)]+\))")


def add_hyperlink(paragraph, url: str, text: str, italic: bool = False) -> None:
    """python-docx has no hyperlink API, so the relationship is added by hand."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    link = docx.oxml.shared.OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    run = docx.oxml.shared.OxmlElement("w:r")
    props = docx.oxml.shared.OxmlElement("w:rPr")
    style = docx.oxml.shared.OxmlElement("w:rStyle")
    style.set(qn("w:val"), "Hyperlink")
    props.append(style)
    if italic:
        props.append(docx.oxml.shared.OxmlElement("w:i"))
    run.append(props)
    label = docx.oxml.shared.OxmlElement("w:t")
    label.text = text
    run.append(label)
    link.append(run)
    paragraph._p.append(link)


def write_runs(paragraph, text: str) -> None:
    """Render `**bold**`, `*italic*` and `[label](url)` into Word runs."""
    for piece in INLINE.split(smart_typography(text)):
        if not piece:
            continue
        if piece.startswith("**") and piece.endswith("**"):
            paragraph.add_run(piece[2:-2]).bold = True
        elif piece.startswith("*") and piece.endswith("*"):
            paragraph.add_run(piece[1:-1]).italic = True
        elif piece.startswith("["):
            label, target = re.match(r"^\[([^\]]+)\]\(([^)]+)\)$", piece).groups()
            if target.startswith(("http://", "https://", "mailto:")):
                add_hyperlink(paragraph, target, label)
            else:
                paragraph.add_run(label)
        else:
            paragraph.add_run(piece)


# --------------------------------------------------------------------------- document


def configure_styles(document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    # Word's default heading blue reads as a template. Black, and sized so the three
    # levels are told apart by weight and space rather than by colour.
    sizes = {"Heading 1": 16, "Heading 2": 13, "Heading 3": 11.5}
    for name, size in sizes.items():
        style = document.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.italic = name == "Heading 3"
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(14 if name == "Heading 1" else 10)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.keep_with_next = True

    caption = document.styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9.5)
    caption.font.italic = True
    caption.font.bold = False
    caption.font.color.rgb = RGBColor(0x40, 0x40, 0x40)


def configure_page(document) -> None:
    for section in document.sections:
        section.page_width, section.page_height = Cm(21), Cm(29.7)
        section.left_margin = section.right_margin = Cm(2.5)
        section.top_margin = section.bottom_margin = Cm(2.5)


def add_field(paragraph, instruction: str, placeholder: str) -> None:
    """Insert a Word field. `dirty` makes Word offer to refresh it on open."""
    begin = docx.oxml.shared.OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    instr = docx.oxml.shared.OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = docx.oxml.shared.OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = docx.oxml.shared.OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    for element in (begin, instr, separate):
        run = paragraph.add_run()
        run._r.append(element)
    paragraph.add_run(placeholder)
    paragraph.add_run()._r.append(end)


def add_page_numbers(document) -> None:
    footer = document.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.style = document.styles["Footer"]
    add_field(footer, " PAGE ", "1")


def render_title_page(document, meta: dict[str, str]) -> None:
    def line(text: str, size: float, bold: bool = False, space: int = 6, italic: bool = False) -> None:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(space)
        run = paragraph.add_run(smart_typography(text))
        run.font.size, run.bold, run.italic = Pt(size), bold, italic

    for _ in range(4):
        document.add_paragraph()
    line(meta.get("title", ""), 20, bold=True, space=10)
    if meta.get("subtitle"):
        line(meta["subtitle"], 14, space=28, italic=True)
    if meta.get("status"):
        line(meta["status"], 12, bold=True, space=28)
    for entry in meta.get("authors", "").splitlines():
        if entry.strip():
            line(entry.strip(), 11.5, space=2)
    if meta.get("date"):
        document.add_paragraph()
        line(meta["date"], 11.5, space=36)

    if meta.get("disclaimer"):
        for entry in meta["disclaimer"].split("\n\n"):
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(smart_typography(entry.replace("\n", " ").strip()))
            run.font.size, run.italic = Pt(9.5), True
            run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)


def render_table(document, lines: list[str]) -> None:
    header = split_table_row(lines[0])
    alignments = column_alignments(lines[1])
    rows = [split_table_row(line) for line in lines[2:]]

    table = document.add_table(rows=1, cols=len(header))
    table.style = document.styles["Table Grid"]
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for index, text in enumerate(header):
        cell = table.rows[0].cells[index]
        cell.paragraphs[0].alignment = ALIGNMENTS[alignments[index] if index < len(alignments) else "left"]
        cell.paragraphs[0].paragraph_format.space_after = Pt(2)
        run = cell.paragraphs[0].add_run(smart_typography(text))
        run.bold = True

    for cells in rows:
        row = table.add_row()
        for index, text in enumerate(cells):
            if index >= len(header):
                break
            paragraph = row.cells[index].paragraphs[0]
            paragraph.alignment = ALIGNMENTS[alignments[index] if index < len(alignments) else "left"]
            paragraph.paragraph_format.space_after = Pt(2)
            write_runs(paragraph, text)

    # Repeat the header row when a table breaks across a page.
    header_props = docx.oxml.shared.OxmlElement("w:trPr")
    header_props.append(docx.oxml.shared.OxmlElement("w:tblHeader"))
    table.rows[0]._tr.insert(0, header_props)


CAPTION = re.compile(r"^(Table|Figure)\s+\d+\.")


def render(document, blocks: list[Block]) -> None:
    for block in blocks:
        if block.kind == "pagebreak":
            document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        elif block.kind == "toc":
            # Formatted to match Heading 1 rather than using the style, so that the
            # contents does not list itself as its own first entry.
            heading = document.add_paragraph()
            heading.paragraph_format.space_before = Pt(0)
            heading.paragraph_format.space_after = Pt(4)
            run = heading.add_run("Contents")
            run.bold, run.font.size = True, Pt(16)
            paragraph = document.add_paragraph()
            add_field(
                paragraph,
                ' TOC \\o "1-2" \\h \\z \\u ',
                "Right-click here and choose Update Field to build the contents.",
            )

        elif block.kind == "heading":
            document.add_paragraph(smart_typography(block.lines[0]), style=f"Heading {min(block.level, 3)}")

        elif block.kind == "table":
            render_table(document, block.lines)

        elif block.kind == "list":
            style = "List Number" if block.ordered else "List Bullet"
            for entry in block.lines:
                paragraph = document.add_paragraph(style=style)
                paragraph.paragraph_format.space_after = Pt(4)
                write_runs(paragraph, entry)

        elif block.kind == "quote":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(1)
            paragraph.paragraph_format.right_indent = Cm(1)
            write_runs(paragraph, " ".join(block.lines))
            for run in paragraph.runs:
                run.italic = True

        else:
            text = " ".join(block.lines)
            style = "Caption" if CAPTION.match(text) else None
            paragraph = document.add_paragraph(style=style)
            if style == "Caption":
                paragraph.add_run(smart_typography(text))
            else:
                write_runs(paragraph, text)


# --------------------------------------------------------------------------- checks


def verify(path: Path, blocks: list[Block]) -> list[str]:
    """Reopen the written file and check the things that break silently."""
    problems: list[str] = []
    document = docx.Document(path)

    styles = {paragraph.style.name for paragraph in document.paragraphs}
    for level in ("Heading 1", "Heading 2"):
        if level not in styles:
            problems.append(f"no paragraph uses the real {level} style")

    for index, table in enumerate(document.tables, start=1):
        if not table.rows:
            problems.append(f"table {index} has no rows")
            continue
        if not any(run.bold for cell in table.rows[0].cells for p in cell.paragraphs for run in p.runs):
            problems.append(f"table {index} has no bold header row")
        widths = {len(row.cells) for row in table.rows}
        if len(widths) > 1:
            problems.append(f"table {index} has ragged rows: {sorted(widths)}")

    expected_tables = sum(1 for block in blocks if block.kind == "table")
    if len(document.tables) != expected_tables:
        problems.append(f"wrote {len(document.tables)} tables, source has {expected_tables}")

    xml = document.element.xml
    if "TOC \\o" not in xml.replace("\\\\", "\\"):
        problems.append("the contents field did not survive into the file")

    # Numbered sections must run 1..N with nothing missing or repeated.
    numbers = [
        int(match.group(1))
        for paragraph in document.paragraphs
        if paragraph.style.name == "Heading 1"
        for match in [re.match(r"^(\d+)\.\s", paragraph.text)]
        if match
    ]
    if numbers and numbers != list(range(1, len(numbers) + 1)):
        problems.append(f"section numbers are {numbers}, expected 1..{len(numbers)}")

    # The author's own proposals were cut from this report by decision, not by accident.
    # These phrases are specific to them and appear nowhere in the Round Table record.
    banned = ("Human-First", "Learning Cycle", "six-day", "Observe reality", "Predict before AI")
    text = "\n".join(p.text for p in document.paragraphs).lower()
    for phrase in banned:
        if phrase.lower() in text:
            problems.append(f"cut material survived: {phrase!r}")

    return problems


def size_report(blocks: list[Block]) -> tuple[int, float]:
    words = 0
    rows = 0
    for block in blocks:
        if block.kind == "table":
            rows += len(block.lines) - 1
            words += sum(len(line.split()) for line in block.lines)
        else:
            words += sum(len(line.split()) for line in block.lines)
    pages = words / WORDS_PER_PAGE + rows / ROWS_PER_PAGE
    return words, pages


def main() -> int:
    if not SOURCE.exists():
        print(f"Missing report source: {SOURCE}")
        return 1

    meta, body = parse_front_matter(SOURCE.read_text(encoding="utf-8"))
    blocks = parse_blocks(body)

    document = docx.Document()
    configure_page(document)
    configure_styles(document)
    document.core_properties.title = meta.get("title", "")
    document.core_properties.subject = meta.get("subtitle", "")
    document.core_properties.comments = meta.get("status", "")

    render_title_page(document, meta)
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    render(document, blocks)
    add_page_numbers(document)

    # python-docx seeds a new document with one empty paragraph before anything else.
    first = document.paragraphs[0]
    if not first.text.strip() and not first.runs:
        first._element.getparent().remove(first._element)

    document.save(OUTPUT)

    words, pages = size_report(blocks)
    print(f"Wrote {OUTPUT.name}")
    print(f"  {words:,} words, {len(document.tables)} tables, ~{pages:.0f} pages estimated")

    problems = verify(OUTPUT, blocks)
    if problems:
        for problem in problems:
            print(f"  - {problem}")
        print(f"Check failed: {len(problems)} problem(s).")
        return 1
    print("  Checks passed: heading styles, table headers, contents field, section numbering")
    return 0


if __name__ == "__main__":
    sys.exit(main())
